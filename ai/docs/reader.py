import os
import re


def read_text_file(path: str) -> str:
    """Lee un archivo de texto plano con utf-8 o fallback a latin-1."""
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(path, "r", encoding="latin-1") as f:
                return f.read()
        except Exception as e:
            return f"[Error leyendo archivo de texto: {str(e)}]"


def read_pdf_file(path: str) -> str:
    """Extrae texto real de un PDF."""
    try:
        import pypdf

        reader = pypdf.PdfReader(path)
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"[PAGINA {i + 1}]\n{page_text}")
        if text_parts:
            return "\n".join(text_parts)
    except ImportError:
        pass
    except Exception as e:
        print(f"Error con pypdf: {e}")

    return read_pdf_fallback(path)


def read_docx_file(path: str) -> str:
    """Extrae texto real de un DOCX."""
    try:
        import docx

        doc = docx.Document(path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text:
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        if text_parts:
            return "\n".join(text_parts)
    except ImportError:
        pass
    except Exception as e:
        print(f"Error con python-docx: {e}")

    return read_docx_fallback(path)


def read_pdf_fallback(path: str) -> str:
    filename = os.path.basename(path)
    try:
        with open(path, "rb") as f:
            data = f.read()

        strings = re.findall(rb"[a-zA-Z0-9\s\.,;:!\?\-\(\)]{10,200}", data)
        if len(strings) > 10:
            extracted = []
            for s in strings[:150]:
                try:
                    extracted.append(s.decode("utf-8", errors="ignore"))
                except Exception:
                    pass
            if len(extracted) > 5:
                return f"[Texto parcial recuperado del binario de {filename}]\n" + "\n".join(extracted)
    except Exception:
        pass

    return (
        f"[Sin texto extraible del PDF: {filename}]\n"
        "Instala pypdf o usa un PDF con capa de texto para indexar contenido real."
    )


def read_docx_fallback(path: str) -> str:
    filename = os.path.basename(path)
    return (
        f"[Sin texto extraible del DOCX: {filename}]\n"
        "Instala python-docx o usa un documento compatible para indexar contenido real."
    )


def extract_text(path: str) -> str:
    """Extrae texto de cualquier formato soportado."""
    if not os.path.exists(path):
        return f"[Error: El archivo no existe en la ruta {path}]"

    ext = os.path.splitext(path)[1].lower()
    if ext == ".txt":
        return read_text_file(path)
    if ext == ".pdf":
        return read_pdf_file(path)
    if ext in [".docx", ".doc"]:
        return read_docx_file(path)

    return read_text_file(path)
